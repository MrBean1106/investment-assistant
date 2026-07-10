"""Render investment reports to PDF (reportlab) and Word (python-docx).

Report content is the JSON stored on the Report model:
    {
      "title": str,
      "sections": [{"heading": str, "content": str}, ...],
      "conclusion": str
    }
Both renderers are CJK-safe (STSong-Light for PDF, system fonts for DOCX).
"""

from io import BytesIO
from typing import Optional

# ── PDF (reportlab) ────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont

# Register a built-in Chinese CID font so CJK text renders without external TTFs.
try:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    _CJK_FONT = "STSong-Light"
except Exception:  # pragma: no cover — fall back to default if unavailable
    _CJK_FONT = "Helvetica"

_PDF_STYLES = {
    "title": ParagraphStyle(
        "Title", fontName=_CJK_FONT, fontSize=18, leading=24,
        spaceAfter=14, textColor=colors.HexColor("#1a1a2e"),
    ),
    "heading": ParagraphStyle(
        "Heading", fontName=_CJK_FONT, fontSize=13, leading=18,
        spaceBefore=12, spaceAfter=6, textColor=colors.HexColor("#2563eb"),
    ),
    "body": ParagraphStyle(
        "Body", fontName=_CJK_FONT, fontSize=10.5, leading=17, spaceAfter=6,
    ),
    "conclusion": ParagraphStyle(
        "Conclusion", fontName=_CJK_FONT, fontSize=10.5, leading=17,
        spaceBefore=10, spaceAfter=6, textColor=colors.HexColor("#374151"),
    ),
}


def _clean(value: Optional[str]) -> str:
    if not value:
        return ""
    # Escape XML special chars for reportlab Paragraph.
    return (str(value).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def render_report_pdf(content: dict) -> bytes:
    title = _clean(content.get("title")) or "招商研判报告"
    sections = content.get("sections") or []
    conclusion = _clean(content.get("conclusion"))

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
        title=title,
    )
    flow = [Paragraph(title, _PDF_STYLES["title"])]

    if isinstance(sections, list):
        for sec in sections:
            if not isinstance(sec, dict):
                continue
            heading = _clean(sec.get("heading"))
            body = _clean(sec.get("content"))
            if heading:
                flow.append(Paragraph(heading, _PDF_STYLES["heading"]))
            if body:
                flow.append(Paragraph(body, _PDF_STYLES["body"]))
            flow.append(Spacer(1, 4))

    if conclusion:
        flow.append(Paragraph("总结建议", _PDF_STYLES["heading"]))
        flow.append(Paragraph(conclusion, _PDF_STYLES["conclusion"]))

    doc.build(flow)
    buf.seek(0)
    return buf.getvalue()


# ── Word (python-docx) ─────────────────────────
def render_report_docx(content: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    title = (content.get("title") or "招商研判报告")
    sections = content.get("sections") or []
    conclusion = content.get("conclusion")

    doc = Document()
    doc.styles["Normal"].font.name = "SimSun"
    doc.styles["Normal"].font.size = Pt(11)

    h = doc.add_heading(title, level=0)
    # Center the title.
    for run in h.runs:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    if isinstance(sections, list):
        for sec in sections:
            if not isinstance(sec, dict):
                continue
            heading = sec.get("heading")
            body = sec.get("content")
            if heading:
                hh = doc.add_heading(heading, level=1)
                for run in hh.runs:
                    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            if body:
                doc.add_paragraph(str(body))

    if conclusion:
        ch = doc.add_heading("总结建议", level=1)
        for run in ch.runs:
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        doc.add_paragraph(str(conclusion))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
